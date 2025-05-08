import json
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox

def parse_har_file(har_file_path):
    """Parse the .har file and extract relevant details."""
    with open(har_file_path, 'r') as file:
        har_data = json.load(file)
    
    entries = har_data['log']['entries']
    request_data = []
    for entry in entries:
        response_time = entry['time']
        
        if response_time <= 2500:
            performance = "Good"  
        elif 2500 < response_time <= 4000:
            performance = "Needs Improvement"  
        else:
            performance = "Poor" 
        
        request_info = {
            "url": entry['request']['url'],
            "method": entry['request']['method'],
            "status": entry['response']['status'],
            "response_time": response_time,
            "response_size": entry['response']['content']['size'],
            "performance": performance
        }
        request_data.append(request_info)
    
    return request_data

def generate_document(request_data, output_doc_path):
    """Generate a Word document from the parsed .har file data."""
    doc = Document()
    doc.add_heading('HAR File Analysis Report', 0)
    
    doc.add_heading('Overview:', level=1)
    doc.add_paragraph('This document contains a detailed analysis of the requests and responses from the provided HAR file.')

    doc.add_heading('Request and Response Details:', level=1)
    table = doc.add_table(rows=1, cols=6)  
    table.style = 'Table Grid'
    
    headers = ["URL", "Method", "Status", "Response Time (ms)", "Response Size (bytes)", "Performance"]
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    
    for request in request_data:
        row = table.add_row().cells
        row[0].text = request['url']
        row[1].text = request['method']
        row[2].text = str(request['status'])
        row[3].text = str(request['response_time'])
        row[4].text = str(request['response_size'])
        row[5].text = request['performance']  

    doc.save(output_doc_path)
    return output_doc_path

def select_har_file():
    """Allow the user to select the .har file."""
    file_path = filedialog.askopenfilename(title="Select HAR File", filetypes=[("HAR Files", "*.har")])
    if file_path:
        har_file_path.set(file_path)

def select_output_location():
    """Allow the user to select the output location for the document."""
    output_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
    if output_path:
        output_doc_path.set(output_path)

def start_conversion():
    """Start the process of converting the HAR file to a document."""
    har_file = har_file_path.get()
    output_file = output_doc_path.get()

    if not har_file or not output_file:
        messagebox.showerror("Error", "Please select both the HAR file and the output location.")
        return
    
    try:
        request_data = parse_har_file(har_file)

        result = generate_document(request_data, output_file)

        messagebox.showinfo("Success", f"Document successfully created at {result}")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

root = tk.Tk()
root.title("HAR File to Document Converter")
root.geometry("400x300")

har_file_path = tk.StringVar()
output_doc_path = tk.StringVar()

tk.Label(root, text="Select HAR File:").pack(pady=10)
tk.Entry(root, textvariable=har_file_path, width=40).pack(pady=5)
tk.Button(root, text="Browse", command=select_har_file).pack(pady=5)

tk.Label(root, text="Select Output Document Location:").pack(pady=10)
tk.Entry(root, textvariable=output_doc_path, width=40).pack(pady=5)
tk.Button(root, text="Browse", command=select_output_location).pack(pady=5)

tk.Button(root, text="Start Conversion", command=start_conversion).pack(pady=20)

root.mainloop()
